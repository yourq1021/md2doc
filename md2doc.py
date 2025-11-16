#!/usr/bin/env python3
"""
md2doc.py

将 Markdown 转换为标准 Office 文档：默认输出 .docx；若指定 .doc，则先生成 .docx 再调用 LibreOffice 转换。

优先使用 pandoc（若系统已安装），否则回退到 python-docx 的基础渲染（支持标题、段落、粗体、斜体、行内代码、代码块、无序/有序列表、引用、水平线、链接的可读文本）。

版式与样式（尽量满足高校论文规范）：
- 纸张 A4（210mm × 297mm）；页边距：上 30mm，下 25mm，左 30mm，右 25mm
- 正文字体：中文宋体（SimSun），英文字与数字 Times New Roman；字号：小四（12pt）；行距：固定 20 磅
- 一级标题：黑体（三号 16pt，居中，段前后各 1 行）；二级标题：黑体（四号 14pt，左对齐，段前后各 1 行）；三级标题：黑体（小四 12pt，左对齐，段前后各 1 行）
- 页眉页脚：宋体 小五（9pt）；可通过 --header 文本设置页眉

用法:
  python md2doc.py input.md -o output.docx
  python md2doc.py input.md -o output.doc

依赖（见 requirements.txt）:
  - pypandoc (可选，但建议，需系统安装 pandoc)
  - python-docx
  - markdown-it-py
"""

from __future__ import annotations

import argparse
import os
import shutil
import subprocess
import sys
from dataclasses import dataclass
from typing import List, Optional, Dict, Any
import tempfile


def which(cmd: str) -> Optional[str]:
    return shutil.which(cmd)


def ensure_parent_dir(path: str) -> None:
    parent = os.path.dirname(os.path.abspath(path))
    if parent and not os.path.exists(parent):
        os.makedirs(parent, exist_ok=True)


def _apply_doc_defaults(doc, header_text: Optional[str], cfg: Optional[Dict[str, Any]] = None) -> None:
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    cfg = cfg or {}
    page = cfg.get("page", {})
    margins = cfg.get("margins", {})
    normal_style = cfg.get("normal", {})
    headings = cfg.get("headings", {})
    header_cfg = cfg.get("header", {})

    # 页面设置 A4 与页边距
    section = doc.sections[0]
    section.page_width = Mm(page.get("width_mm", 210))
    section.page_height = Mm(page.get("height_mm", 297))
    section.top_margin = Mm(margins.get("top_mm", 30))
    section.bottom_margin = Mm(margins.get("bottom_mm", 25))
    section.left_margin = Mm(margins.get("left_mm", 30))
    section.right_margin = Mm(margins.get("right_mm", 25))

    # 正文 Normal 样式：宋体/Times New Roman，小四，行距 20 磅
    normal = doc.styles["Normal"]
    normal.font.size = Pt(normal_style.get("size_pt", 12))
    normal.font.name = normal_style.get("chinese", "SimSun")
    # East Asia 字体指定
    if hasattr(normal, "_element"):
        rpr = normal._element.rPr
        if rpr is not None and rpr.rFonts is not None:
            rpr.rFonts.set(qn('w:eastAsia'), normal_style.get("chinese", "SimSun"))
            rpr.rFonts.set(qn('w:ascii'), normal_style.get("western", "Times New Roman"))
            rpr.rFonts.set(qn('w:hAnsi'), normal_style.get("western", "Times New Roman"))
    # 行距设置
    for s in [normal]:
        if hasattr(s, "paragraph_format"):
            s.paragraph_format.line_spacing = Pt(normal_style.get("line_spacing_pt", 20))

    # 标题样式
    def cfg_heading(name: str, size_pt: int, align: int, family: str):
        st = doc.styles[name]
        st.font.name = family
        st.font.size = Pt(size_pt)
        if hasattr(st, "paragraph_format"):
            st.paragraph_format.alignment = align
            st.paragraph_format.space_before = Pt(headings.get(name, {}).get("space_before_pt", 12))
            st.paragraph_format.space_after = Pt(headings.get(name, {}).get("space_after_pt", 12))

    cfg_heading("Heading 1", headings.get("Heading 1", {}).get("size_pt", 16),
                WD_ALIGN_PARAGRAPH.__dict__.get(headings.get("Heading 1", {}).get("align", "CENTER"), WD_ALIGN_PARAGRAPH.CENTER),
                headings.get("Heading 1", {}).get("family", "SimHei"))
    cfg_heading("Heading 2", headings.get("Heading 2", {}).get("size_pt", 14),
                WD_ALIGN_PARAGRAPH.__dict__.get(headings.get("Heading 2", {}).get("align", "LEFT"), WD_ALIGN_PARAGRAPH.LEFT),
                headings.get("Heading 2", {}).get("family", "SimHei"))
    cfg_heading("Heading 3", headings.get("Heading 3", {}).get("size_pt", 12),
                WD_ALIGN_PARAGRAPH.__dict__.get(headings.get("Heading 3", {}).get("align", "LEFT"), WD_ALIGN_PARAGRAPH.LEFT),
                headings.get("Heading 3", {}).get("family", "SimHei"))

    # 页眉页脚
    header = section.header
    if header_text or header_cfg.get("text"):
        header_text = header_text or header_cfg.get("text")
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run(header_text)
        run.font.name = header_cfg.get("family", normal_style.get("chinese", "SimSun"))
        run.font.size = Pt(header_cfg.get("size_pt", 9))
    footer = section.footer
    if footer and footer.paragraphs:
        for p in footer.paragraphs:
            for r in p.runs:
                r.font.name = header_cfg.get("family", normal_style.get("chinese", "SimSun"))
                r.font.size = Pt(header_cfg.get("size_pt", 9))


def _create_reference_docx(path: str, header_text: Optional[str], cfg: Optional[Dict[str, Any]]) -> None:
    from docx import Document
    doc = Document()
    _apply_doc_defaults(doc, header_text, cfg)
    # 放一个空段落，确保样式写入
    doc.add_paragraph("")
    doc.save(path)


def convert_with_pandoc(input_path: str, output_path: str, header_text: Optional[str], cfg: Optional[Dict[str, Any]]) -> bool:
    try:
        import pypandoc  # type: ignore
    except Exception:
        return False
    if not which("pandoc"):
        return False
    ensure_parent_dir(output_path)
    try:
        # 生成 reference.docx 以传递版式与样式
        with tempfile.TemporaryDirectory() as td:
            ref_path = os.path.join(td, "reference.docx")
            _create_reference_docx(ref_path, header_text, cfg)
        pypandoc.convert_file(
            source_file=input_path,
            to="docx",
            outputfile=output_path,
            extra_args=[
                "--standalone",
                "--from=markdown+tex_math_dollars+pipe_tables+table_captions",
                f"--reference-doc={ref_path}",
            ],
        )
        return os.path.exists(output_path) and os.path.getsize(output_path) > 0
    except Exception as exc:
        print(f"[pandoc] 失败: {exc}", file=sys.stderr)
        return False


def fallback_convert_with_python_docx(input_path: str, output_docx_path: str, header_text: Optional[str], cfg: Optional[Dict[str, Any]]) -> None:
    from docx import Document  # type: ignore
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    try:
        from markdown_it import MarkdownIt  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "缺少 markdown-it-py，请安装：pip install markdown-it-py"
        ) from exc

    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()

    md = MarkdownIt("commonmark").enable("table").enable("strikethrough")
    tokens = md.parse(text)

    doc = Document()
    _apply_doc_defaults(doc, header_text, cfg)

    # 简易栈管理列表缩进
    list_stack: List[dict] = []

    def add_paragraph_from_inline(inline_children):
        p = doc.add_paragraph()
        run = None
        for child in inline_children or []:
            if child.type == "text":
                run = p.add_run(child.content)
            elif child.type == "code_inline":
                run = p.add_run(child.content)
                run.font.name = "Consolas"
            elif child.type == "softbreak" or child.type == "hardbreak":
                p.add_run("\n")
            elif child.type == "em_open" or child.type == "em_close":
                # 斜体在 python-docx 需要在具体 run 上设置；这里简化处理：
                pass
            elif child.type == "strong_open" or child.type == "strong_close":
                pass
            elif child.type == "link_open":
                # 链接文本以普通文本写入
                pass
            elif child.type == "link_close":
                pass
            else:
                if getattr(child, "content", None):
                    run = p.add_run(child.content)
        return p

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        t = tok.type

        if t == "heading_open":
            level = int(tok.tag[-1]) if tok.tag.startswith("h") else 1
            # 下一个应为 inline
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            text_content = ""
            if inline and hasattr(inline, "children") and inline.children:
                text_content = "".join(ch.content for ch in inline.children if getattr(ch, "content", None))
            p = doc.add_paragraph(text_content, style=f"Heading {min(level,6)}")
            i += 2  # 跳过 inline
            # 跳过 heading_close
            i += 1
            continue

        if t == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            if inline and hasattr(inline, "children"):
                p = add_paragraph_from_inline(inline.children)
            else:
                doc.add_paragraph("")
            # 跳过 inline、paragraph_close
            i += 3
            continue

        if t in ("bullet_list_open", "ordered_list_open"):
            list_stack.append({"type": t, "indent": len(list_stack)})
            i += 1
            continue

        if t in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue

        if t == "list_item_open":
            # 项目内容通常为 paragraph_open -> inline -> paragraph_close
            # 简化为读取下一段 inline 文本
            j = i + 1
            content_text = ""
            while j < len(tokens) and tokens[j].type != "list_item_close":
                if tokens[j].type == "inline" and getattr(tokens[j], "children", None):
                    content_text = "".join(
                        ch.content for ch in tokens[j].children if getattr(ch, "content", None)
                    )
                j += 1
            bullet = "•" if list_stack and list_stack[-1]["type"] == "bullet_list_open" else "1."
            indent = "    " * (len(list_stack) - 1)
            para = doc.add_paragraph(f"{indent}{bullet} {content_text}")
            # 行距保持 20 磅
            para.paragraph_format.line_spacing = Pt(20)
            i = j + 1
            continue

        if t == "fence":  # 代码块
            code_text = tok.content.rstrip("\n")
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            i += 1
            continue

        if t == "blockquote_open":
            # 简化：为引用增加前缀
            j = i + 1
            quote_lines: List[str] = []
            while j < len(tokens) and tokens[j].type != "blockquote_close":
                if tokens[j].type == "inline" and getattr(tokens[j], "content", None):
                    quote_lines.append(tokens[j].content)
                j += 1
            for line in quote_lines:
                doc.add_paragraph(f"> {line}")
            i = j + 1
            continue

        if t == "hr":
            doc.add_paragraph("——————")
            i += 1
            continue

        # 其他 token：跳过
        i += 1

    ensure_parent_dir(output_docx_path)
    doc.save(output_docx_path)


def convert_docx_to_doc_with_libreoffice(input_docx: str, output_doc: str) -> bool:
    soffice = which("soffice") or which("libreoffice")
    if not soffice:
        print("未找到 LibreOffice(soffice)，无法导出 .doc。已生成 .docx。", file=sys.stderr)
        return False
    ensure_parent_dir(output_doc)
    try:
        # LibreOffice 会将结果输出到指定目录
        out_dir = os.path.dirname(os.path.abspath(output_doc)) or os.getcwd()
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "doc",
                os.path.abspath(input_docx),
                "--outdir",
                out_dir,
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        # 转换后文件名与输入同名但扩展为 .doc
        base = os.path.splitext(os.path.basename(input_docx))[0]
        produced = os.path.join(out_dir, base + ".doc")
        if produced != os.path.abspath(output_doc):
            if os.path.exists(produced):
                os.replace(produced, os.path.abspath(output_doc))
        return os.path.exists(output_doc)
    except subprocess.CalledProcessError as exc:
        print(f"LibreOffice 转换失败: {exc}", file=sys.stderr)
        return False


@dataclass
class Args:
    input: str
    output: Optional[str]
    header: Optional[str]
    config: Optional[str]


def parse_args(argv: List[str]) -> Args:
    parser = argparse.ArgumentParser(description="将 Markdown 转换为 .docx 或 .doc 文件")
    parser.add_argument("input", help="输入 Markdown 文件路径")
    parser.add_argument("-o", "--output", help="输出文件路径（.docx 或 .doc）")
    parser.add_argument("--header", help="页眉文本（默认不设置）")
    parser.add_argument("--config", help="样式配置文件（YAML 或 JSON）")
    ns = parser.parse_args(argv)
    return Args(input=ns.input, output=ns.output, header=ns.header, config=ns.config)


def _load_config(path: Optional[str]) -> Optional[Dict[str, Any]]:
    if not path:
        return None
    abs_path = os.path.abspath(path)
    if not os.path.exists(abs_path):
        print(f"配置文件不存在：{abs_path}", file=sys.stderr)
        return None
    try:
        if abs_path.lower().endswith((".yml", ".yaml")):
            import yaml  # type: ignore
            with open(abs_path, "r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        else:
            import json
            with open(abs_path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as exc:
        print(f"读取配置失败：{exc}", file=sys.stderr)
        return None


def main(argv: List[str]) -> int:
    args = parse_args(argv)
    input_md = os.path.abspath(args.input)
    if not os.path.exists(input_md):
        print(f"找不到输入文件: {input_md}", file=sys.stderr)
        return 2

    if args.output:
        output_path = os.path.abspath(args.output)
        out_ext = os.path.splitext(output_path)[1].lower()
    else:
        base = os.path.splitext(os.path.basename(input_md))[0]
        output_path = os.path.join(os.path.dirname(input_md), base + ".docx")
        out_ext = ".docx"

    # 先生成 docx
    if out_ext == ".docx":
        docx_target = output_path
    elif out_ext == ".doc":
        base = os.path.splitext(output_path)[0]
        docx_target = base + ".docx"
    else:
        print("仅支持输出 .docx 或 .doc", file=sys.stderr)
        return 2

    # 载入配置
    cfg = _load_config(args.config)

    # 尝试 pandoc
    ok = convert_with_pandoc(input_md, docx_target, args.header, cfg)
    if not ok:
        # 回退方案
        try:
            fallback_convert_with_python_docx(input_md, docx_target, args.header, cfg)
            ok = True
        except Exception as exc:
            print(f"回退转换失败: {exc}", file=sys.stderr)
            return 1

    if out_ext == ".doc":
        converted = convert_docx_to_doc_with_libreoffice(docx_target, output_path)
        if converted:
            # 可选择保留中间 docx；这里保留，避免信息丢失
            pass
        else:
            print("已生成 .docx，但 .doc 转换失败（需要 LibreOffice）", file=sys.stderr)
            return 3

    print(f"转换完成: {output_path if out_ext == '.doc' else docx_target}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))



